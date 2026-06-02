from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('RewardHub - User Manual', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Version info
version = doc.add_paragraph('Version 1.0.0 | June 2, 2026 | Production Ready ✅')
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
version_format = version.runs[0]
version_format.font.size = Pt(10)
version_format.font.italic = True

doc.add_paragraph()

# What is RewardHub
doc.add_heading('What is RewardHub?', 1)
doc.add_paragraph('RewardHub is a cashback rewards platform that tracks user purchases and automatically calculates cashback rewards. It features a web dashboard for managing users and transactions, plus a REST API for integration.')

# Getting Started
doc.add_heading('Getting Started (5 Minutes)', 1)
doc.add_heading('Prerequisites', 2)
doc.add_paragraph('Java 17 or higher')
doc.add_paragraph('Maven 3.8 or higher')

doc.add_heading('Installation', 2)
doc.add_paragraph('Step 1: Clone and navigate', style='List Number')
doc.add_paragraph('git clone https://github.com/nayanpunmiya/RewardHub.git', style='List Bullet')
doc.add_paragraph('cd RewardHub', style='List Bullet')

doc.add_paragraph('Step 2: Build', style='List Number')
doc.add_paragraph('mvn clean package', style='List Bullet')

doc.add_paragraph('Step 3: Run', style='List Number')
doc.add_paragraph('java -jar target/reward-hub-1.0.0.jar', style='List Bullet')

doc.add_paragraph('Step 4: Open browser', style='List Number')
doc.add_paragraph('http://localhost:8080', style='List Bullet')

# Frontend Dashboard
doc.add_heading('Frontend Dashboard - What You See', 1)
doc.add_heading('Home Page', 2)
doc.add_paragraph('When you open the application, you\'ll see:', style='List Bullet')
p = doc.add_paragraph('Total Users - Number of registered users')
p = doc.add_paragraph('Total Transactions - Number of completed transactions')
p = doc.add_paragraph('Total Cashback - Total amount distributed as cashback')

# How to Use
doc.add_heading('How to Use the Frontend', 1)

doc.add_heading('Creating a User', 2)
table = doc.add_table(rows=2, cols=2)
table.style = 'Light Grid Accent 1'
table.cell(0, 0).text = 'Where'
table.cell(0, 1).text = 'Click "Users" → "Create New User"'
table.cell(1, 0).text = 'What to Enter'
table.cell(1, 1).text = 'Name (required)\nEmail (required)'

doc.add_heading('Viewing Users', 2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
table.cell(0, 0).text = 'Where'
table.cell(0, 1).text = 'Click "Users" → "View All Users"'

doc.add_heading('Creating a Transaction', 2)
table = doc.add_table(rows=2, cols=2)
table.style = 'Light Grid Accent 1'
table.cell(0, 0).text = 'What to Enter'
table.cell(0, 1).text = 'User ID\nAmount (e.g., 100.50)\nDescription (e.g., "Coffee")\nCashback % (e.g., 5)'
table.cell(1, 0).text = 'What Happens'
table.cell(1, 1).text = 'Transaction is recorded\nCashback is calculated automatically\nUser\'s total cashback is updated'

doc.add_paragraph('Example: Purchase $100 with 5% cashback = $5 earned')

# API Endpoints
doc.add_heading('API Endpoints - For Developers', 1)

doc.add_heading('Health Check', 2)
doc.add_paragraph('URL: http://localhost:8080/api/health', style='List Bullet')
doc.add_paragraph('Method: GET', style='List Bullet')
doc.add_paragraph('Response: {"status":"OK","message":"API is running"}', style='List Bullet')

doc.add_heading('Get Statistics', 2)
doc.add_paragraph('URL: http://localhost:8080/api/stats', style='List Bullet')
doc.add_paragraph('Method: GET', style='List Bullet')

doc.add_heading('Create User', 2)
doc.add_paragraph('URL: POST http://localhost:8080/api/users', style='List Bullet')
doc.add_paragraph('Send: {"name":"John","email":"john@example.com"}', style='List Bullet')

doc.add_heading('Get All Users', 2)
doc.add_paragraph('URL: GET http://localhost:8080/api/users', style='List Bullet')

doc.add_heading('Create Transaction', 2)
doc.add_paragraph('URL: POST http://localhost:8080/api/transactions', style='List Bullet')
doc.add_paragraph('Send: {"userId":"USER1","amount":100,"description":"Shopping","cashbackPercentage":5}', style='List Bullet')

# Features
doc.add_heading('Features Overview', 1)
doc.add_paragraph('Create new users', style='List Bullet')
doc.add_paragraph('View all users', style='List Bullet')
doc.add_paragraph('Update user information', style='List Bullet')
doc.add_paragraph('Create transactions', style='List Bullet')
doc.add_paragraph('Automatic cashback calculation', style='List Bullet')
doc.add_paragraph('View statistics in real-time', style='List Bullet')
doc.add_paragraph('REST API access', style='List Bullet')

# Common Workflows
doc.add_heading('Common Workflows', 1)

doc.add_heading('Workflow 1: Set Up a New Customer', 2)
doc.add_paragraph('Open Dashboard', style='List Number')
doc.add_paragraph('Click "Users"', style='List Number')
doc.add_paragraph('Click "Create New User"', style='List Number')
doc.add_paragraph('Enter name and email', style='List Number')
doc.add_paragraph('Click "Create"', style='List Number')

doc.add_heading('Workflow 2: Record a Purchase', 2)
doc.add_paragraph('Click "Transactions"', style='List Number')
doc.add_paragraph('Click "Create New Transaction"', style='List Number')
doc.add_paragraph('Select user from dropdown', style='List Number')
doc.add_paragraph('Enter amount, description, and cashback %', style='List Number')
doc.add_paragraph('Click "Create"', style='List Number')

# Troubleshooting
doc.add_heading('Troubleshooting', 1)

doc.add_heading('Application won\'t start', 2)
doc.add_paragraph('Error: "Address already in use :8080"', style='List Bullet')
doc.add_paragraph('Solution: Use different port with --server.port=9000', style='List Bullet')

doc.add_heading('Can\'t find Java', 2)
doc.add_paragraph('Error: "java: command not found"', style='List Bullet')
doc.add_paragraph('Solution: Install Java 17+ from https://www.java.com/', style='List Bullet')

# System Requirements
doc.add_heading('System Requirements', 1)
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = 'Requirement'
header_cells[1].text = 'Minimum'
header_cells[2].text = 'Recommended'

table.cell(1, 0).text = 'Java'
table.cell(1, 1).text = '17'
table.cell(1, 2).text = '17+'

table.cell(2, 0).text = 'Maven'
table.cell(2, 1).text = '3.8'
table.cell(2, 2).text = '3.9+'

table.cell(3, 0).text = 'RAM'
table.cell(3, 1).text = '256 MB'
table.cell(3, 2).text = '512 MB+'

# Key Points
doc.add_heading('Key Points to Remember', 1)
doc.add_paragraph('No database setup required - Works out of the box', style='List Bullet')
doc.add_paragraph('Data is in-memory - Resets on application restart', style='List Bullet')
doc.add_paragraph('REST API included - Can integrate with other apps', style='List Bullet')
doc.add_paragraph('Complete dashboard - No need to use API for basic tasks', style='List Bullet')
doc.add_paragraph('Cashback auto-calculated - No manual calculation needed', style='List Bullet')
doc.add_paragraph('Production-ready - Can be deployed to cloud', style='List Bullet')

# Next Steps
doc.add_heading('Next Steps', 1)
doc.add_paragraph('Run the application (follow Getting Started section)', style='List Number')
doc.add_paragraph('Create test users (use Frontend Dashboard)', style='List Number')
doc.add_paragraph('Create test transactions (use Frontend Dashboard)', style='List Number')
doc.add_paragraph('Explore the API (use cURL or Postman)', style='List Number')
doc.add_paragraph('Deploy to cloud (optional)', style='List Number')

# Save
doc.save('MANUAL.docx')
print("✅ MANUAL.docx created successfully!")
